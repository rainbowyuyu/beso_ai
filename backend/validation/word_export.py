"""Export validation report as Word (.docx) with embedded figures."""
from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from backend.validation.report_builder import CATEGORY_LABELS
from backend.validation.scorer import ValidationScore
from backend.validation.word_export_common import (
    FIGURE_CAPTIONS,
    METRIC_LABELS,
    add_bullet,
    add_heading,
    add_paragraph,
    add_table,
    embed_figures,
    format_metric,
    load_clause_map,
    payload_from_dir,
    payload_from_score,
    require_docx,
)

# Re-export for backward compatibility
__all__ = ["FIGURE_CAPTIONS", "METRIC_LABELS", "build_validation_docx"]


def _build_docx_body(
    doc,
    payload: dict[str, Any],
    *,
    validation_id: str,
    geometry_title: str,
    out_dir: Path,
    llm_rationales: dict[str, str] | None = None,
) -> None:
    clauses = load_clause_map()
    metrics = payload.get("metrics") or {}
    bc = payload.get("benchmark_context") or {}

    add_heading(doc, f"验证报告 — {geometry_title}", level=1)
    add_paragraph(doc, f"验证 ID：{validation_id}")
    add_paragraph(
        doc,
        f"生成时间：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
    )
    doc.add_paragraph()

    add_heading(doc, "执行摘要", level=2)
    add_bullet(doc, f"综合得分：{payload.get('overall_score', 0):.1f} / 100")
    add_bullet(doc, f"等级：{payload.get('grade', '—')}")
    add_bullet(doc, f"钢耗强度：{metrics.get('steel_intensity_t_per_MW', 0):.1f} t/MW")
    add_bullet(doc, f"估算总钢量：{metrics.get('steel_mass_t_est', 0):.0f} t")
    add_bullet(doc, f"估算方法：{metrics.get('steel_mass_t_source', 'n/a')}")

    if bc:
        doc.add_paragraph()
        add_heading(doc, "基准对比", level=3)
        add_bullet(doc, f"20 MW 机队分位：{bc.get('percentile_vs_fleet_20mw', '—')}%")
        add_bullet(doc, f"机队中位钢耗：{bc.get('fleet_median_intensity_20mw', '—')} t/MW")
        add_bullet(doc, f"相对图强：{bc.get('delta_vs_tuqiang_pct', '—')}%")

    doc.add_paragraph()
    add_heading(doc, "AI Review 五维得分", level=2)
    ai_scores = payload.get("ai_review_scores") or {}
    ai_labels = payload.get("ai_review_labels") or {}
    ai_metrics = payload.get("ai_review_metrics") or {}
    ai_rows: list[list[str]] = []
    for key, val in ai_scores.items():
        label = ai_labels.get(key, key)
        measured = ai_metrics.get(key)
        mtxt = f"{measured:.2f}" if isinstance(measured, (int, float)) else "—"
        ai_rows.append([label, mtxt, f"{float(val):.1f}"])
    add_table(doc, ["指标", "实测", "得分"], ai_rows)

    add_heading(doc, "合规维度得分（DNV 规则明细）", level=2)
    cat_rows = [
        [CATEGORY_LABELS.get(cat, cat), f"{float(val):.1f}"]
        for cat, val in (payload.get("category_scores") or {}).items()
    ]
    add_table(doc, ["维度", "得分"], cat_rows)

    add_heading(doc, "验证图表", level=2)
    embed_figures(doc, out_dir)

    add_heading(doc, "指标实测", level=2)
    metric_rows: list[list[str]] = []
    for key, label in METRIC_LABELS.items():
        v = metrics.get(key)
        if v is not None:
            metric_rows.append([label, format_metric(v)])
    add_table(doc, ["指标", "数值"], metric_rows)

    add_heading(doc, "规则明细", level=2)
    rule_rows: list[list[str]] = []
    for r in payload.get("rules") or []:
        measured = f"{r['measured']:.3f}" if r.get("measured") is not None else "—"
        reg = str(r.get("regulation_ref") or r.get("source") or "")[:48]
        rule_rows.append([
            r.get("id", ""),
            r.get("status", ""),
            f"{float(r.get('score_0_100', 0)):.0f}",
            measured,
            str(r.get("threshold", "")),
            reg,
        ])
    add_table(doc, ["规则", "状态", "得分", "实测", "阈值", "法规/来源"], rule_rows)

    add_heading(doc, "法规条款映射", level=2)
    clause_rows: list[list[str]] = []
    for r in payload.get("rules") or []:
        cref = r.get("clause_ref") or ""
        c = clauses.get(cref, {})
        summary = str(c.get("summary_zh") or "—")[:80]
        doc_name = c.get("doc") or cref
        clause_rows.append([r.get("id", ""), str(doc_name), summary])
    add_table(doc, ["规则", "条款", "摘要"], clause_rows)

    if llm_rationales:
        add_heading(doc, "LLM 合规解释（不参与打分）", level=2)
        for rid, text in llm_rationales.items():
            add_heading(doc, rid, level=3)
            add_paragraph(doc, text.strip())

    notes = payload.get("calibration_notes") or []
    if notes:
        add_heading(doc, "综合分校准说明", level=2)
        for note in notes:
            add_bullet(doc, str(note))

    assumptions = payload.get("assumptions") or []
    if assumptions:
        add_heading(doc, "假设与局限", level=2)
        for a in assumptions:
            add_bullet(doc, str(a))

    doc.add_paragraph()
    add_paragraph(
        doc,
        "本报告由 AI Engineer 验证模块自动生成；法规符合性需经持证验船师复核。",
    )


def build_validation_docx(
    out_dir: Path,
    *,
    score: ValidationScore | None = None,
    validation_id: str | None = None,
    geometry_title: str | None = None,
    llm_rationales: dict[str, str] | None = None,
) -> Path:
    """Build validation_report.docx under out_dir. Requires python-docx."""
    require_docx()
    from docx import Document

    out_dir = out_dir.resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    vid = validation_id or out_dir.name

    if score is not None:
        payload = payload_from_score(score, validation_id=vid)
        title = geometry_title or "Design candidate"
    else:
        payload, title = payload_from_dir(out_dir)
        if geometry_title:
            title = geometry_title

    docx_path = out_dir / "validation_report.docx"
    doc = Document()
    _build_docx_body(
        doc,
        payload,
        validation_id=vid,
        geometry_title=title,
        out_dir=out_dir,
        llm_rationales=llm_rationales,
    )
    doc.save(str(docx_path))
    return docx_path
