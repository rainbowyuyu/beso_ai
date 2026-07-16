"""Shared helpers for validation Word export (simple and detailed)."""
from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import yaml

from backend.validation.paths import FIGURE_STEMS
from backend.validation.report_builder import score_to_json
from backend.validation.scorer import ValidationScore

_REPO_ROOT = Path(__file__).resolve().parents[2]
CLAUSE_INDEX = _REPO_ROOT / "rules" / "dnv_clause_index.yaml"
VALIDATION_RULES = _REPO_ROOT / "rules" / "validation_rules.yaml"

FIGURE_CAPTIONS: dict[str, str] = {
    "fig_benchmark_position": "图1 钢耗强度 · 行业基准位置（含行业趋势线）",
    "fig_benchmark_capacity": "图2 单机容量 · 行业基准位置（含行业趋势线）",
    "fig_benchmark_unit_cost": "图3 单位造价 · 行业基准位置（含行业趋势线）",
    "fig_benchmark_construction": "图4 施工年限 · 行业基准位置（含行业趋势线）",
    "fig_benchmark_fatigue": "图5 疲劳寿命 · 行业基准位置（含行业趋势线）",
    "fig_score_radar": "图6 AI Review 五维对比雷达（机队对比）",
    "fig_fleet_metrics_bars": "图7 AI Review 有效性（原始指标 vs AI/法规得分）",
    "fig_rule_heatmap": "图8 规则得分分解",
    "fig_capacity_intensity": "图9 容量–钢耗强度散点",
    "fig_ai_review_validity": "图10 同一五维指标：AI Review vs 法规标准打分对照",
}

METRIC_LABELS: dict[str, str] = {
    "target_power_MW": "目标容量 (MW)",
    "steel_intensity_t_per_MW": "钢耗强度 (t/MW)",
    "steel_mass_t_est": "估算钢量 (t)",
    "unit_cost_cny_per_MW": "单位造价 (万元/MW)",
    "construction_years": "施工年限 (年)",
    "fatigue_life_years": "疲劳寿命 (年)",
    "draft_m": "吃水 (m)",
    "wall_thickness_m": "壳体壁厚 (m)",
    "leg_mean_diameter_m": "柱均径 (m)",
    "leg_mean_length_m": "柱均长 (m)",
    "leg_slenderness_L_over_D": "长细比 L/D",
    "leg_diameter_uniformity_pct": "柱径一致性 (%)",
    "leg_taper_ratio": "变径比",
    "column_spacing_m": "柱间距 (m)",
    "beso3_column_spacing_m": "BESO3 柱间距 (m)",
    "design_domain_span_m": "设计域尺度 (m)",
    "leg_layout_angle_deg_std": "120°布局偏差 (°)",
    "top_plate_diameter_m": "顶盘直径 (m)",
    "plate_to_leg_diameter_ratio": "顶盘/柱径比",
    "dt_ratio": "D/t 比",
    "hub_elevation_m": "Hub 高程 (m)",
    "freeboard_proxy_m": "干舷代理 (m)",
    "bbox_z_m": "竖向包络 (m)",
    "scale_factor": "缩放系数",
    "steel_mass_t_source": "钢量估算方法",
}


def require_docx():
    try:
        from docx import Document  # noqa: F401
    except ImportError as e:
        raise ImportError(
            "python-docx is required for Word export; "
            "install with: pip install -r backend/requirements-validation.txt"
        ) from e


def load_clause_map() -> dict[str, dict[str, Any]]:
    if not CLAUSE_INDEX.is_file():
        return {}
    raw = yaml.safe_load(CLAUSE_INDEX.read_text(encoding="utf-8"))
    return {c["id"]: c for c in (raw.get("clauses") or [])}


def load_validation_rules() -> dict[str, Any]:
    if not VALIDATION_RULES.is_file():
        return {}
    return yaml.safe_load(VALIDATION_RULES.read_text(encoding="utf-8")) or {}


def geometry_title_from_dir(out_dir: Path, fallback: str = "Design candidate") -> str:
    snap = out_dir / "input_geometry_snapshot.json"
    if snap.is_file():
        try:
            data = json.loads(snap.read_text(encoding="utf-8"))
            title = str(data.get("title") or "").strip()
            if title:
                return title
        except (OSError, ValueError, json.JSONDecodeError):
            pass
    return fallback


def load_geometry_snapshot(out_dir: Path) -> dict[str, Any]:
    snap = out_dir / "input_geometry_snapshot.json"
    if not snap.is_file():
        return {}
    try:
        return json.loads(snap.read_text(encoding="utf-8"))
    except (OSError, ValueError, json.JSONDecodeError):
        return {}


def payload_from_score(score: ValidationScore, *, validation_id: str) -> dict[str, Any]:
    return score_to_json(score, validation_id=validation_id, artifacts={})


def payload_from_dir(out_dir: Path) -> tuple[dict[str, Any], str]:
    score_path = out_dir / "validation_score.json"
    if not score_path.is_file():
        raise FileNotFoundError(f"Missing {score_path}")
    payload = json.loads(score_path.read_text(encoding="utf-8"))
    title = geometry_title_from_dir(out_dir)
    return payload, title


def set_run_cn(run, *, size_pt: float | None = 11, bold: bool = False, italic: bool = False) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_cn(run, size_pt=16 if level == 1 else (14 if level == 2 else 12), bold=True)


def add_paragraph(doc, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_cn(run, bold=bold, italic=italic)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_run_cn(run)


def add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                set_run_cn(run, bold=True)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    set_run_cn(run)
    doc.add_paragraph()


def format_metric(value: Any) -> str:
    if value is None:
        return "—"
    if isinstance(value, float):
        return f"{value:.3f}"
    return str(value)


def setup_page_margins(doc) -> None:
    from docx.shared import Cm

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def add_page_break(doc) -> None:
    doc.add_page_break()


def add_toc(doc, *, placeholder: str = "目录（请在 Word 中右键「更新域」）") -> None:
    """Insert a TOC field; Word updates page numbers on open."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def embed_figures(doc, out_dir: Path, *, figure_width_in: float = 6.2) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    for stem in FIGURE_STEMS:
        png = out_dir / f"{stem}.png"
        if not png.is_file():
            continue
        caption = FIGURE_CAPTIONS.get(stem, stem)
        try:
            doc.add_picture(str(png), width=Inches(figure_width_in))
        except Exception:
            add_paragraph(doc, f"[无法嵌入图片：{png.name}]")
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        set_run_cn(run, size_pt=10)
        doc.add_paragraph()
