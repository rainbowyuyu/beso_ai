"""Export detailed design-basis-style validation report as Word (.docx)."""
from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from backend.validation.report_builder import CATEGORY_LABELS
from backend.validation.scorer import ValidationScore
from backend.validation.word_export_common import (
    METRIC_LABELS,
    add_bullet,
    add_heading,
    add_page_break,
    add_paragraph,
    add_table,
    add_toc,
    embed_figures,
    format_metric,
    load_clause_map,
    load_geometry_snapshot,
    load_validation_rules,
    payload_from_dir,
    payload_from_score,
    require_docx,
    set_run_cn,
    setup_page_margins,
)


def _add_cover(doc, *, geometry_title: str, validation_id: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("AI Review 设计评审详细报告")
    set_run_cn(run, size_pt=22, bold=True)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run(geometry_title)
    set_run_cn(run2, size_pt=16, bold=True)

    doc.add_paragraph()
    for line in (
        f"验证 ID：{validation_id}",
        f"生成时间：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
        "AI Engineer · 设计验证模块",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_run_cn(run, size_pt=11)

    add_page_break(doc)


def _section_1_overview(
    doc,
    *,
    geometry_title: str,
    payload: dict[str, Any],
) -> None:
    rules_cfg = load_validation_rules()
    clauses = load_clause_map()

    add_heading(doc, "1 项目概况", level=1)

    add_heading(doc, "1.1 工程简介", level=2)
    add_paragraph(
        doc,
        f"本报告针对候选漂浮式海上风电基础方案「{geometry_title}」开展 AI Review 设计评审。"
        "评审在详细设计审图前，对容量、材料效率、经济性、建造周期与服役寿命等多目标进行可量化比选，"
        "输出五维综合得分及与行业已建成/规划机队的对标结论。",
    )

    add_heading(doc, "1.2 设计方案说明", level=2)
    metrics = payload.get("metrics") or {}
    add_paragraph(
        doc,
        f"目标单机容量 {metrics.get('target_power_MW', '—')} MW；"
        f"估算钢耗强度 {metrics.get('steel_intensity_t_per_MW', 0):.1f} t/MW；"
        f"估算总钢量约 {metrics.get('steel_mass_t_est', 0):.0f} t。"
        "几何参数由参数化重建与 BESO 拓扑优化结果提取，详见第 2 章。",
    )

    add_heading(doc, "1.3 规范与标准", level=2)
    add_paragraph(doc, "本评审规则库主要引用以下规范与行业文件：")
    std_rows: list[list[str]] = []
    seen_docs: set[str] = set()
    for c in clauses.values():
        doc_name = str(c.get("doc") or "")
        if not doc_name or doc_name in seen_docs:
            continue
        seen_docs.add(doc_name)
        std_rows.append([doc_name, str(c.get("summary_zh") or "—")[:60]])
    if not std_rows:
        std_rows = [
            ["DNV-ST-0437", "漂浮式海上风电结构设计"],
            ["DNVGL-OS-C301", "稳性与水密性"],
            ["DNVGL-RP-0286", "疲劳评估"],
        ]
    add_table(doc, ["规范/文件", "适用范围"], std_rows)

    add_heading(doc, "1.4 AI Review 方法摘要", level=2)
    ai_cfg = rules_cfg.get("ai_review") or {}
    weights = ai_cfg.get("dimension_weights") or {}
    labels = ai_cfg.get("dimension_labels_zh") or payload.get("ai_review_labels") or {}
    weight_parts = [
        f"{labels.get(k, k)} {float(v) * 100:.0f}%"
        for k, v in weights.items()
    ]
    add_paragraph(
        doc,
        "AI Review 将候选体映射为五项可观测工程指标（容量、钢耗、造价、施工年限、疲劳寿命），"
        "经归一化子得分与加权综合分输出方案排序；"
        "并行构建法规评审通道，在相同五维指标上依据船级社规范与行业阈值独立打分，"
        "通过机队对照检验两通道一致性。",
    )
    if weight_parts:
        add_paragraph(doc, f"五维权重：{'；'.join(weight_parts)}。")

    add_heading(doc, "1.5 报告范围说明", level=2)
    add_paragraph(
        doc,
        "本报告为 AI Review 自动化评审产物，内容限于几何提取指标、五维打分、机队对标与规则校核结果。"
        "不包含场址环境（风/浪/流/水位/地质）、详细荷载工况组合、有限元应力校核、系泊系统详算等"
        "需专项分析或外部输入的内容；上述章节不出现在本报告中。"
        "法规符合性结论需经持证验船师复核后方可用于审图。",
    )


def _section_2_geometry(doc, *, payload: dict[str, Any], geometry: dict[str, Any]) -> None:
    metrics = payload.get("metrics") or {}

    add_heading(doc, "2 设计输入与几何参数", level=1)

    add_heading(doc, "2.1 目标容量与经济性指标", level=2)
    econ_rows = [
        ["目标容量", format_metric(metrics.get("target_power_MW")), "MW"],
        ["钢耗强度", format_metric(metrics.get("steel_intensity_t_per_MW")), "t/MW"],
        ["估算总钢量", format_metric(metrics.get("steel_mass_t_est")), "t"],
        ["单位造价", format_metric(metrics.get("unit_cost_cny_per_MW")), "万元/MW"],
        ["施工年限", format_metric(metrics.get("construction_years")), "年"],
        ["疲劳寿命", format_metric(metrics.get("fatigue_life_years")), "年"],
    ]
    add_table(doc, ["参数", "数值", "单位"], econ_rows)

    add_heading(doc, "2.2 浮体几何与结构布局", level=2)
    layout_rows = [
        [label, format_metric(metrics.get(key))]
        for key, label in METRIC_LABELS.items()
        if key not in {
            "target_power_MW",
            "steel_intensity_t_per_MW",
            "steel_mass_t_est",
            "unit_cost_cny_per_MW",
            "construction_years",
            "fatigue_life_years",
            "steel_mass_t_source",
        }
        and metrics.get(key) is not None
    ]
    add_table(doc, ["参数", "数值"], layout_rows)

    legs = (
        geometry.get("beso7_method1_topology_reconstructed", {}).get("legs")
        or geometry.get("legs")
        or []
    )
    if legs:
        add_heading(doc, "2.3 立柱分段参数", level=2)
        leg_rows: list[list[str]] = []
        for leg in legs:
            leg_rows.append([
                str(leg.get("id") or leg.get("name") or "—"),
                f"{float(leg.get('diameter_mm', 0)) / 1000:.3f}" if leg.get("diameter_mm") else "—",
                f"{float(leg.get('length_mm', 0)) / 1000:.2f}" if leg.get("length_mm") else "—",
                f"{float(leg.get('radius_mm', 0)) / 1000:.3f}" if leg.get("radius_mm") else "—",
            ])
        add_table(doc, ["柱号", "直径 (m)", "长度 (m)", "半径 (m)"], leg_rows)

    assumptions = payload.get("assumptions") or []
    if assumptions:
        add_heading(doc, "2.4 假设与数据来源", level=2)
        for a in assumptions:
            add_bullet(doc, str(a))


def _validity_rows_for_cohort(
    cohort: list[dict[str, Any]],
    ai_cols: list[dict[str, str]],
    reg_cols: list[dict[str, str]],
) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in cohort:
        raw = row.get("raw_metrics") or {}
        ai = row.get("ai_scores") or {}
        reg = row.get("regulatory_scores") or {}
        cells = [str(row.get("name") or "—")]
        for c in ai_cols:
            v = raw.get(c["key"])
            cells.append(str(v) if v is not None else "—")
        for c in ai_cols:
            v = ai.get(c["key"])
            cells.append(f"{v:.1f}" if v is not None else "—")
        for c in reg_cols:
            v = reg.get(c["key"])
            cells.append(f"{v:.1f}" if v is not None else "—")
        rows.append(cells)
    return rows


def _add_validity_table(
    doc,
    validity: dict[str, Any],
    *,
    title: str,
    cohort: list[dict[str, Any]],
) -> None:
    ai_cols = validity.get("ai_columns") or []
    reg_cols = validity.get("regulatory_columns") or []
    if not ai_cols or not cohort:
        return

    add_heading(doc, title, level=3)
    raw_hdr = [c.get("label", c["key"]) for c in ai_cols]
    ai_hdr = [f"AI·{c.get('label', c['key'])}" for c in ai_cols]
    reg_hdr = [f"法规·{c.get('label', c['key'])}" for c in reg_cols]
    headers = ["项目"] + raw_hdr + ai_hdr + reg_hdr
    rows = _validity_rows_for_cohort(cohort, ai_cols, reg_cols)
    if rows:
        add_table(doc, headers, rows)


def _section_3_ai_review(doc, *, payload: dict[str, Any], out_dir: Path) -> None:
    add_heading(doc, "3 AI Review 评分与机队对标", level=1)

    add_heading(doc, "3.1 五维指标与 AI 得分", level=2)
    ai_scores = payload.get("ai_review_scores") or {}
    ai_labels = payload.get("ai_review_labels") or {}
    ai_metrics = payload.get("ai_review_metrics") or {}
    ai_weights = payload.get("ai_review_weights") or {}
    ai_rows: list[list[str]] = []
    for key, val in ai_scores.items():
        label = ai_labels.get(key, key)
        measured = ai_metrics.get(key)
        mtxt = f"{measured:.2f}" if isinstance(measured, (int, float)) else "—"
        wt = ai_weights.get(key)
        wtxt = f"{float(wt) * 100:.0f}%" if wt is not None else "—"
        ai_rows.append([label, mtxt, f"{float(val):.1f}", wtxt])
    add_table(doc, ["指标", "实测", "AI 得分", "权重"], ai_rows)

    add_heading(doc, "3.2 法规通道五维得分", level=2)
    reg_scores = payload.get("regulatory_review_scores") or {}
    reg_rows: list[list[str]] = []
    for key in ai_scores:
        label = ai_labels.get(key, key)
        measured = ai_metrics.get(key)
        mtxt = f"{measured:.2f}" if isinstance(measured, (int, float)) else "—"
        rval = reg_scores.get(key)
        rtxt = f"{float(rval):.1f}" if rval is not None else "—"
        reg_rows.append([label, mtxt, rtxt])
    add_table(doc, ["指标", "实测", "法规得分"], reg_rows)
    add_paragraph(
        doc,
        f"法规综合分：{payload.get('regulatory_overall', '—')} / 100",
        bold=True,
    )

    bc = payload.get("benchmark_context") or {}
    if bc:
        add_heading(doc, "3.3 基准对比", level=2)
        bench_rows = [
            ["20 MW 机队分位", f"{bc.get('percentile_vs_fleet_20mw', '—')}", "%"],
            ["机队中位钢耗", str(bc.get("fleet_median_intensity_20mw", "—")), "t/MW"],
            ["相对图强偏差", f"{bc.get('delta_vs_tuqiang_pct', '—')}", "%"],
            ["AI Review 综合分", f"{payload.get('overall_score', 0):.1f}", "/100"],
            ["等级", str(payload.get("grade", "—")), ""],
        ]
        add_table(doc, ["指标", "数值", "单位"], bench_rows)

    validity = payload.get("validity_table") or {}
    if validity:
        add_heading(doc, "3.4 AI Review 有效性（机队对照）", level=2)
        note = validity.get("note") or ""
        if note:
            add_paragraph(doc, note)
        vs = validity.get("validity_summary") or {}
        if vs:
            summary_rows = [
                ["样本量 n", str(vs.get("n", "—"))],
                ["综合分 Spearman ρ", str(vs.get("overall_spearman", "—"))],
                ["五维平均 |AI−法规|", str(vs.get("overall_mean_abs_diff", "—"))],
                ["高一致占比", f"{vs.get('high_agreement_pct', '—')}%"],
            ]
            add_table(doc, ["统计量", "数值"], summary_rows)

        _add_validity_table(
            doc,
            validity,
            title="表3-1 · 已建成机队",
            cohort=validity.get("commissioned_cohort") or [],
        )
        _add_validity_table(
            doc,
            validity,
            title="表3-2 · 规划/前瞻项目",
            cohort=validity.get("planned_cohort") or [],
        )
        cand = validity.get("candidate")
        if cand:
            _add_validity_table(
                doc,
                validity,
                title="表3-3 · 本方案（候选）",
                cohort=[cand],
            )

    add_heading(doc, "3.5 对标图表", level=2)
    embed_figures(doc, out_dir)


def _section_4_compliance(
    doc,
    *,
    payload: dict[str, Any],
    llm_rationales: dict[str, str] | None = None,
) -> None:
    clauses = load_clause_map()

    add_heading(doc, "4 合规规则校核", level=1)

    add_heading(doc, "4.1 维度得分汇总", level=2)
    cat_rows = [
        [CATEGORY_LABELS.get(cat, cat), f"{float(val):.1f}"]
        for cat, val in (payload.get("category_scores") or {}).items()
    ]
    add_table(doc, ["维度", "得分"], cat_rows)

    add_heading(doc, "4.2 规则明细", level=2)
    rule_rows: list[list[str]] = []
    for r in payload.get("rules") or []:
        measured = f"{r['measured']:.3f}" if r.get("measured") is not None else "—"
        desc = str(r.get("description_zh") or "")[:40]
        rule_rows.append([
            r.get("id", ""),
            CATEGORY_LABELS.get(r.get("category", ""), r.get("category", "")),
            r.get("status", ""),
            f"{float(r.get('score_0_100', 0)):.0f}",
            measured,
            str(r.get("threshold", "")),
            desc,
        ])
    add_table(doc, ["规则", "维度", "状态", "得分", "实测", "阈值", "说明"], rule_rows)

    add_heading(doc, "4.3 法规条款映射", level=2)
    clause_rows: list[list[str]] = []
    for r in payload.get("rules") or []:
        cref = r.get("clause_ref") or ""
        c = clauses.get(cref, {})
        summary = str(c.get("summary_zh") or "—")[:80]
        doc_name = c.get("doc") or cref
        clause_rows.append([
            r.get("id", ""),
            str(doc_name),
            str(r.get("regulation_ref") or r.get("source") or "")[:40],
            summary,
        ])
    add_table(doc, ["规则", "条款文件", "法规引用", "摘要"], clause_rows)

    if llm_rationales:
        add_heading(doc, "4.4 LLM 合规解释（不参与打分）", level=2)
        for rid, text in llm_rationales.items():
            add_heading(doc, rid, level=3)
            add_paragraph(doc, text.strip())

    notes = payload.get("calibration_notes") or []
    if notes:
        add_heading(doc, "4.5 综合分校准说明", level=2)
        for note in notes:
            add_bullet(doc, str(note))


def _build_detailed_body(
    doc,
    payload: dict[str, Any],
    *,
    validation_id: str,
    geometry_title: str,
    out_dir: Path,
    geometry: dict[str, Any],
    llm_rationales: dict[str, str] | None = None,
) -> None:
    _add_cover(doc, geometry_title=geometry_title, validation_id=validation_id)

    add_heading(doc, "目录", level=1)
    add_toc(doc)
    add_page_break(doc)

    _section_1_overview(doc, geometry_title=geometry_title, payload=payload)
    _section_2_geometry(doc, payload=payload, geometry=geometry)
    _section_3_ai_review(doc, payload=payload, out_dir=out_dir)
    _section_4_compliance(doc, payload=payload, llm_rationales=llm_rationales)

    doc.add_paragraph()
    add_paragraph(
        doc,
        "— 报告结束 —",
        bold=True,
    )
    add_paragraph(
        doc,
        "本报告由 AI Engineer 验证模块自动生成；法规符合性需经持证验船师复核。",
    )


def build_validation_docx_detailed(
    out_dir: Path,
    *,
    score: ValidationScore | None = None,
    validation_id: str | None = None,
    geometry_title: str | None = None,
    llm_rationales: dict[str, str] | None = None,
) -> Path:
    """Build validation_report_detailed.docx under out_dir. Requires python-docx."""
    require_docx()
    from docx import Document

    out_dir = out_dir.resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    vid = validation_id or out_dir.name

    if score is not None:
        payload = payload_from_score(score, validation_id=vid)
        title = geometry_title or "Design candidate"
    else:
        payload, title = payload_from_dir(out_dir)
        if geometry_title:
            title = geometry_title

    geometry = load_geometry_snapshot(out_dir)

    docx_path = out_dir / "validation_report_detailed.docx"
    doc = Document()
    setup_page_margins(doc)
    _build_detailed_body(
        doc,
        payload,
        validation_id=vid,
        geometry_title=title,
        out_dir=out_dir,
        geometry=geometry,
        llm_rationales=llm_rationales,
    )
    doc.save(str(docx_path))
    return docx_path
