#!/usr/bin/env python3
"""Two-sentence Nature-style supplement: automated CAD drawing before certification."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
OUT = REPO / "docs" / "AI_Engineer_CAD_Drawing_Nature_Supplement.docx"


def set_run_en(run, *, size: float = 11) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def set_run_cn(run, *, size: float = 10, color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.italic = True
    if color:
        run.font.color.rgb = color


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    note = doc.add_paragraph()
    set_run_cn(
        note.add_run("【作者用】插入 Fig. 2 Step 8 或尺寸优化段落后，替换原 third phase 出图描述。"),
        size=9,
        color=RGBColor(0x66, 0x66, 0x66),
    )
    doc.add_paragraph()

    for text in (
        "After dimension optimization and limit-state verification, The AI Engineer projects the "
        "converged parametric solid into standard plan, elevation, and section sheets through the "
        "CAD kernel, automatically populating dimensions and annotations from the same geometric "
        "record that cleared the upstream checks.",
        "The exported general-arrangement drawing package is archived with the verified model and "
        "verification logs as the geometric front matter of the certification dossier, before "
        "automated multidimensional review.",
    ):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        set_run_en(p.add_run(text))

    doc.add_paragraph()
    zh = doc.add_paragraph()
    set_run_cn(
        zh.add_run(
            "中文：尺寸优化与极限工况校核完成后，系统将收敛的参数化实体经 CAD 内核投影为"
            "标准平/立/剖视图，并从同一几何记录自动填充尺寸与标注。"
            "出图文件与校核通过的模型一并归档，作为认证文件集的几何附件，再进入 AI Review。"
        ),
        size=10,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
