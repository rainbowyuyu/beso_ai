#!/usr/bin/env python3
"""Nature-style prose supplement: how AI Engineer plots and writes reports (no code names)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
OUT = REPO / "docs" / "AI_Engineer_PhaseIII_Report_Plotting_Supplement.docx"


def set_run_en(run, *, size: float = 11, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_run_cn(run, *, size: float = 10, italic: bool = True, color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_en(run, size=14 if level == 1 else 12, bold=True)


def add_body(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_en(run)


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ── Author note (not for submission) ─────────────────────────────────────
    note_p = doc.add_paragraph()
    run = note_p.add_run(
        "【作者用 · 非正文】插入位置：AI Engineer-202607042.docx · "
        "「Generating FOWT foundation」小节 · 紧接尺寸优化与 Zwind 校核段落后，"
        "替换或扩展原 “The third phase of The AI Scientist …” 段落。"
        "下文英文块可直接粘贴入稿；不含软件模块名，仅流程叙述，体例对齐 Nature 正文。"
    )
    set_run_cn(run, size=9, italic=True, color=RGBColor(0x66, 0x66, 0x66))
    doc.add_paragraph()

    add_heading(doc, "Supplementary block — automated plotting and report writing", level=1)

    # ── Main prose (Nature-style, one section) ───────────────────────────────
    add_body(
        doc,
        "The third phase of The AI Engineer converts the numerical outcomes of sizing "
        "optimization and limit-state verification into auditable design documentation "
        "suitable for concept-level submission. This phase follows the same closed-loop "
        "discipline as upstream design: the orchestration layer decides what to render and "
        "in what order, but every figure pixel and every tabulated score is produced by "
        "deterministic backends that read the finalized geometry descriptors, structural "
        "proxies, and fleet-comparison statistics emitted at the end of Phase III. The "
        "language model therefore supervises framing, cross-referencing, and scope "
        "declarations; it does not freehand sketches or invent numerical entries.",
    )

    add_body(
        doc,
        "Visualization proceeds in two coupled steps. First, each candidate is mapped onto "
        "five quantifiable engineering dimensions—rated capacity, steel intensity per "
        "megawatt, unit capital cost, construction period, and fatigue design life—and "
        "compared against a curated fleet of existing and planned floating foundations. "
        "Second, the system renders a fixed family of journal-ready figures from this "
        "shared record: benchmark-position charts that locate the candidate among industry "
        "peers; radar overlays that show dimensional strengths and weaknesses across the "
        "fleet; heatmaps that decompose compliance scores by rule category; and "
        "side-by-side tables that contrast AI-oriented review scores with "
        "classification-society thresholds on the same five metrics. Because all plots "
        "draw from one validation record, any figure reproduced in the report can be "
        "traced back to the same physics-validated experiment that cleared the preceding "
        "optimization gates.",
    )

    add_body(
        doc,
        "Report assembly mirrors conventional offshore foundation practice while remaining "
        "machine-auditable. A concise manuscript layer summarizes design intent, extracted "
        "physical quantities, dimensional sub-scores, and an overall review grade; "
        "structured appendices carry rule-level pass or warning statuses, clause references "
        "to DNV and related standards, and fleet-wide effectiveness statistics that "
        "quantify agreement between the AI review channel and the regulatory channel. "
        "Where a detailed design-basis dossier is required, the same record expands into a "
        "multi-chapter document—with cover, table of contents, design inputs, review "
        "scores, and embedded figures—while explicitly scoping out site-specific "
        "environmental tables, load-case combinations, and high-fidelity stress fields "
        "not yet computed by the agent. To situate the candidate within prior art, the "
        "orchestrator may query existing project reports over multiple reasoning rounds, "
        "informing how comparative baselines and normative references are cited in the "
        "related-work passages without altering the underlying measurements.",
    )

    add_body(
        doc,
        "The deliverables of this phase—a scored review dossier, an embedded figure set, "
        "and exportable manuscripts in human- and machine-readable formats—feed directly "
        "into Phase IV, where The Automated Reviewer evaluates whether the documented "
        "design merits retention within the autonomous search loop or formal submission "
        "to a classification society.",
    )

    doc.add_paragraph()
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_en(sep.add_run("— — —"), size=10, italic=True)

    # ── Chinese mirror (optional paste for 中文审稿沟通) ─────────────────────
    add_heading(doc, "中文对照（可选 · 非投稿正文）", level=2)
    add_body(
        doc,
        "AI Engineer 的第三阶段将尺寸优化与极限工况校核的数值成果转化为可审计的概念设计文件。"
        "编排层决定渲染内容与顺序，但每一张图的像素与表格中的分值均由确定性后端读取 Phase III 末输出的"
        "几何描述符、结构代理量与机队对照统计而生成；语言模型负责叙述框架、交叉引用与范围声明，"
        "而非自由「作画」或编造数值。",
    )
    add_body(
        doc,
        "可视化分两步耦合进行：先将候选方案映射为容量、单位兆瓦用钢、单位造价、施工年限与疲劳寿命五维指标，"
        "并与 curated 机队基准对照；再据此统一渲染基准位置图、机队雷达叠加图、合规规则热力图，"
        "以及 AI 评审通道与法规通道的同维对照表。全部图件共享同一验证记录，保证报告中的任意图表"
        "均可追溯至通过上游优化闸门的同一物理实验。",
    )
    add_body(
        doc,
        "报告组装在保留近海基础设计惯例的同时保持机器可审计：摘要层概括设计意图、提取物理量、"
        "分项得分与综合等级；结构化附录给出规则级通过/警告状态、DNV 等条款引用，"
        "以及 AI 与法规双通道一致性的 fleet 统计。若需详细设计基础体例，同一记录可扩展为"
        "含封面、目录、设计输入、评审得分与嵌入图件的多章文件，并以范围声明明确尚未自动计算的"
        "场址环境、荷载组合与高保真应力场。编排层可经多轮推理查询既有项目报告，"
        "以规范相关工作的引用方式，而不改动底层测量值。该阶段产物直接供 Phase IV 的 Automated Reviewer "
        "判定方案是否保留于自主搜索环或提交船级社。",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path} ({path.stat().st_size} bytes)")
