#!/usr/bin/env python3
"""Generate standalone Word supplement: how AI Engineer writes reports and produces plots."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

REPO = Path(__file__).resolve().parents[1]
OUT = REPO / "docs" / "AI_Engineer_Report_and_Plotting_Supplement.docx"


def set_run_cn(run, *, size: float = 11, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_cn(run, size=16 if level == 1 else (14 if level == 2 else 12), bold=True)


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_cn(run, bold=bold, italic=italic)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_run_cn(run)


def add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                set_run_cn(run, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    set_run_cn(run)
    doc.add_paragraph()


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ── Cover ────────────────────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("AI Engineer 自动化写报告与画图")
    set_run_cn(run, size=20, bold=True)
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_cn(sub_p.add_run("实现说明 · 论文 Phase III–IV 补充稿"), size=12)
    doc.add_page_break()

    # ── 0. 定位 ──────────────────────────────────────────────────────────────
    add_heading(doc, "0. 在 AI Engineer 全流程中的位置", level=1)
    add_para(
        doc,
        "在论文 Fig. 2 所示工作流中，Phase II（b–f）完成拓扑优化与参数化重建，"
        "Phase III（g）完成 SQP/PSO 尺寸优化及 Zwind 极限工况校核。"
        "此后系统进入「出图 + 写报告」阶段：一方面将数值结果可视化为可嵌入设计文件的图表，"
        "另一方面按工程设计报告体例组织文字、表格与图注，形成可审阅、可导出的设计说明与 AI Review 报告。",
    )
    add_para(
        doc,
        "与人类工程师「Drawing and writing report」相对应，AI Engineer 将这一环节拆为两条确定性管线："
        "（1）matplotlib 程序化制图；（2）结构化报告组装（Markdown / JSON / Word）。"
        "LLM 编排层负责触发管线、传递几何 JSON 与候选标签，"
        "但图表像素与报告表格内容均由可复现的 Python 模块生成，而非由模型直接「画出来」或「编出来」。",
    )

    add_heading(doc, "0.1 与前序阶段的衔接", level=2)
    add_bullet(doc, "输入：参数化几何 JSON（含立柱分段、顶盘、设计域尺度等）+ 可选 validation_overrides。")
    add_bullet(doc, "中间量：geometry_metrics 提取 20+ 项物理量；scorer 输出五维 AI 分与 26 条 DNV 规则明细。")
    add_bullet(doc, "机队上下文：fleet_scoring 加载 11 台白名单基准；validity_eval 计算机队 AI–法规一致性。")
    add_bullet(doc, "输出：runs/_validation/<id>/ 目录下的 PNG/PDF 图、MD/JSON 报告及 Word 文档。")

    # ── 1. 总体架构 ──────────────────────────────────────────────────────────
    add_heading(doc, "1. 总体架构：先画图、后写报告", level=1)
    add_para(
        doc,
        "核心入口为 backend/validation/pipeline.py 中的 run_validation()。"
        "该函数按固定顺序串联「度量 → 打分 → 机队对照 → 画图 → 写报告」，"
        "保证同一验证 ID 下图表与文字引用同一套 validation_score.json。",
    )
    add_table(
        doc,
        ["步骤", "模块", "产物"],
        [
            ["1", "geometry_metrics.extract_geometry_metrics", "钢量、柱径、间距等 metrics"],
            ["2", "scorer.score_design", "ValidationScore（五维 + 规则）"],
            ["3", "validity_eval.compute_validity_table", "机队 AI vs 法规对照表"],
            ["4", "plots.generate_all_plots", "10 张 Nature 风格 PNG/PDF"],
            ["5", "report_builder.write_reports", "validation_report.md / .json"],
            ["6", "word_export / word_export_detailed", "简版与详细 Word"],
        ],
    )

    # ── 2. 画图 ──────────────────────────────────────────────────────────────
    add_heading(doc, "2. 画图实现（How to plot）", level=1)
    add_heading(doc, "2.1 技术栈与风格约定", level=2)
    add_bullet(doc, "渲染后端：matplotlib Agg（无 GUI，可在 uvicorn / FreeCAD 子线程中运行）。")
    add_bullet(doc, "配色：NATURE_COLORS 字典（国际蓝 #3C5488、国内红 #E64B35、候选绿 #00A087）。")
    add_bullet(doc, "字体：英文轴标签与图例；论文级图另见 graph/plot_ai_vs_tuqiang_bars.py（八指标 AI vs 图强）。")
    add_bullet(doc, "输出：每张图同时保存 .png（Word 嵌入）与 .pdf（论文排版）。")

    add_heading(doc, "2.2 十类验证图表（plots.py）", level=2)
    add_para(doc, "generate_all_plots() 依次调用下列绘图函数，stem 名与文件一一对应：")
    add_table(
        doc,
        ["图件 stem", "函数", "内容"],
        [
            ["fig_benchmark_position", "plot_benchmark_position", "钢耗强度 · 机队基准位置 + 趋势线"],
            ["fig_benchmark_capacity", "plot_benchmark_metric(capacity)", "单机容量基准位置"],
            ["fig_benchmark_unit_cost", "plot_benchmark_metric(cost)", "单位造价基准位置"],
            ["fig_benchmark_construction", "plot_benchmark_metric(construction)", "施工年限基准位置"],
            ["fig_benchmark_fatigue", "plot_benchmark_metric(fatigue)", "疲劳寿命基准位置"],
            ["fig_score_radar", "plot_score_radar", "11 台机队五维雷达叠加 + 底部表格"],
            ["fig_fleet_metrics_bars", "plot_fleet_metrics_bars", "原始指标 vs AI/法规得分柱状图"],
            ["fig_rule_heatmap", "plot_rule_heatmap", "26 条规则得分热力图"],
            ["fig_capacity_intensity", "plot_capacity_intensity", "容量–钢耗强度散点"],
            ["fig_ai_review_validity", "plot_validity_table", "AI Review vs 法规五维对照表（图形式）"],
        ],
    )

    add_heading(doc, "2.3 基准位置图绘制逻辑", level=2)
    add_para(
        doc,
        "plot_benchmark_metric() 从 benchmark_loader 读取机队记录（11 台白名单），"
        "横轴为项目简称，纵轴为待比指标（如 t/MW）。"
        "候选方案以绿色高亮标注；已建成/国内/国际项目分色。"
        "对钢耗强度子图，可选绘制 300 t/MW 参考虚线及线性趋势线，"
        "用于直观展示候选在 20 MW 样本中的分位与相对图强偏差。",
    )

    add_heading(doc, "2.4 雷达图与有效性图", level=2)
    add_para(
        doc,
        "plot_score_radar() 将 fleet_radar_series 返回的五维归一化序列绘制为极坐标折线，"
        "11 台机队全部叠加，项目名自动缩短以防溢出。"
        "plot_validity_table() 将 validity_table 中的已建成/规划/本方案三组数据"
        "渲染为带色块的表格图，底部附 Spearman ρ、平均 |AI−法规| 等有效性摘要。",
    )

    add_heading(doc, "2.5 独立八指标对比图（graph/）", level=2)
    add_para(
        doc,
        "除验证模块内置十图外，graph/plot_ai_vs_tuqiang_bars.py 读取 "
        "rules/ai_vs_tuqiang_comparison.yaml，生成 AI vs 图强八指标分组柱状图（英文版式，罗马数字 I–VIII 子图编号）。"
        "该脚本与验证管线解耦，供论文 Fig. 3 及宣传材料单独调用：python graph/plot_ai_vs_tuqiang_bars.py。",
    )

    # ── 3. 写报告 ──────────────────────────────────────────────────────────────
    add_heading(doc, "3. 写报告实现（How to write report）", level=1)
    add_heading(doc, "3.1 三层报告产物", level=2)
    add_table(
        doc,
        ["格式", "生成模块", "用途"],
        [
            ["Markdown", "report_builder.build_markdown_report", "Web 预览、版本 diff、论文 Methods 素材"],
            ["JSON", "report_builder.score_to_json", "机器可读全量快照（含 validity_table）"],
            ["Word 简版", "word_export.build_validation_docx", "执行摘要 + 图表嵌入（快速审阅）"],
            ["Word 详细版", "word_export_detailed.build_validation_docx_detailed", "设计基础风格四章结构"],
        ],
    )

    add_heading(doc, "3.2 Markdown 报告结构", level=2)
    add_para(doc, "build_markdown_report() 按固定章节组装纯文本与 Markdown 表格：")
    for item in [
        "执行摘要（综合分、等级、钢耗、基准对比）",
        "AI Review 五维得分表（实测 + 得分）",
        "合规维度得分（benchmark / C301 / ST-0437 / RP-0286）",
        "规则明细与法规条款映射（dnv_clause_index.yaml）",
        "表1 · 机队 AI vs 法规五维对照（三组 cohort）",
        "可选 LLM 条款解释（不参与打分）",
        "综合分校准说明与假设局限",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.3 Word 详细报告（设计基础风格）", level=2)
    add_para(
        doc,
        "word_export_detailed.py 参考图强号设计基础报告体例，但仅写入 AI Review 已有数据：",
    )
    for item in [
        "封面 + 目录（Word TOC 域）",
        "第 1 章 项目概况（工程简介、规范列表、AI Review 方法、报告范围声明）",
        "第 2 章 设计输入与几何参数（经济性表、布局表、立柱分段）",
        "第 3 章 AI Review 评分与机队对标（五维表、有效性表、嵌入全部 PNG）",
        "第 4 章 合规规则校核（维度得分、26 条规则、条款映射）",
    ]:
        add_bullet(doc, item)
    add_para(
        doc,
        "场址环境、FEA 应力云图、系泊详算等 AI Review 未计算的内容不出现在目录中，"
        "在 §1.5 以范围说明交代边界。",
    )

    add_heading(doc, "3.4 共享排版工具", level=2)
    add_para(
        doc,
        "word_export_common.py 提供中英混排（Times New Roman + 宋体）、Table Grid 表格、"
        "TOC 域插入、6.2 inch 宽图嵌入及 FIGURE_CAPTIONS 图注字典，简版与详细版共用。",
    )

    # ── 4. API 与触发 ──────────────────────────────────────────────────────────
    add_heading(doc, "4. API、前端与触发方式", level=1)
    add_table(
        doc,
        ["端点 / 入口", "行为"],
        [
            ["POST /api/validation/run", "上传几何 JSON，跑完整管线并预生成报告"],
            ["GET /api/validation/{id}/report", "返回 Markdown 全文"],
            ["GET /api/validation/{id}/export/word", "下载简版 validation_report.docx"],
            ["GET /api/validation/{id}/export/word/detailed", "下载详细 validation_report_detailed.docx"],
            ["frontend_static/validation.html", "「导出 Word 报告」与「详细设计基础报告」双按钮"],
            ["scripts/verify_validation_module.py", "冒烟：rules/optimized_geometry.json → 全产物"],
        ],
    )

    # ── 5. 与论文叙述的对应 ──────────────────────────────────────────────────
    add_heading(doc, "5. 与论文 Phase III 叙述的衔接（可插入正文）", level=1)
    add_para(
        doc,
        "原稿中「The third phase … produces a concise write-up … using its notes and plots (Methods)」"
        "可补充为如下中英文表述，使「怎么画图、怎么写报告」前后逻辑通顺：",
        italic=True,
    )
    doc.add_paragraph()
    add_para(
        doc,
        "【英文】After sizing optimization and Zwind limit-state verification (Phase III), "
        "The AI Engineer invokes a deterministic reporting pipeline rather than prompting the LLM to "
        "draft figures freehand. Structured geometry JSON and scored metrics feed "
        "matplotlib-based Nature-style plots (fleet benchmark positions, five-dimensional radar, "
        "rule heatmaps, and AI–regulatory validity tables). These artifacts, together with "
        "clause-traceable rule results, are assembled into Markdown, JSON, and Word reports "
        "via report_builder and word_export modules. A detailed design-basis-style Word export "
        "mirrors conventional offshore foundation reports (cover, table of contents, four chapters) "
        "while explicitly scoping out site-specific environmental tables not yet computed by the agent. "
        "The LLM orchestrator triggers and audits this pipeline; all numerical tables and pixel outputs "
        "remain reproducible from validation_score.json and the SHA-256 artifact bundle.",
    )
    doc.add_paragraph()
    add_para(
        doc,
        "【中文】尺寸优化与 Zwind 极限工况校核（Phase III）完成后，"
        "AI Engineer 调用确定性报告管线，而非由大模型自由「作画」或「撰稿」。"
        "结构化的几何 JSON 与打分结果依次驱动 matplotlib 程序化制图"
        "（机队基准位置图、五维雷达图、规则热力图、AI–法规有效性对照表等十类图件），"
        "并与可追溯的 DNV 规则明细一并写入 Markdown、JSON 及 Word 报告"
        "（report_builder / word_export 模块）。"
        "详细 Word 导出采用工程设计基础报告体例（封面、目录、四章），"
        "同时以范围声明明确尚未自动计算的场址环境等内容。"
        "LLM 编排层负责触发与审计该管线；所有表格数值与图像像素均可由 "
        "validation_score.json 及 SHA-256 工件包复现。",
    )

    add_heading(doc, "5.1 建议插入位置", level=2)
    add_bullet(doc, "论文正文：「Generating FOWT foundation」小节 Phase III 段落后。")
    add_bullet(doc, "Methods：「Automated design review module」段前，作为 Phase III→IV 过渡段。")
    add_bullet(doc, "Fig. 2 图注：步骤 (h) 可注明「08_AI_Review_Report 由 validation 管线自动生成」。")

    # ── 6. 复现命令 ──────────────────────────────────────────────────────────
    add_heading(doc, "6. 本地复现", level=1)
    add_para(doc, "pip install -r backend/requirements-validation.txt", bold=True)
    add_para(doc, "python scripts/verify_validation_module.py", bold=True)
    add_para(doc, "python -m pytest tests/test_validation_scorer.py -v", bold=True)
    add_para(doc, "python graph/plot_ai_vs_tuqiang_bars.py", bold=True)
    add_para(
        doc,
        "产物目录：runs/_validation/_smoke_latest/；"
        "含 validation_report.docx、validation_report_detailed.docx 及 fig_*.png。",
    )

    doc.add_paragraph()
    add_para(
        doc,
        "— 本文档由 scripts/generate_ai_engineer_report_plotting_docx.py 自动生成；"
        "与仓库 backend/validation/ 实现保持同步。",
        italic=True,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path} ({path.stat().st_size} bytes)")
