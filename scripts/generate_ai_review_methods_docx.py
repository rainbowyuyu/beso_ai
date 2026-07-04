#!/usr/bin/env python3
"""Generate AI Review Methods section as Word manuscript (Nature-style, paper prose)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

REPO = Path(__file__).resolve().parents[1]
OUT = REPO / "docs" / "AI_Review_Methods_Manuscript.docx"


def set_run_cn(run, *, size: float = 11, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_cn(run, size=14 if level == 1 else 12, bold=True)


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_cn(run, bold=bold, italic=italic)


def add_equation(doc, text: str, label: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_cn(run, italic=True)
    if label:
        run2 = p.add_run(f"    ({label})")
        set_run_cn(run2)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_run_cn(run)


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ── Methods ──────────────────────────────────────────────────────────────
    add_heading(doc, "Methods", level=1)

    add_para(
        doc,
        "漂浮式海上风电基础在详细设计审图前，需要在容量、材料效率、经济性、建造周期与服役寿命"
        "等多目标之间进行可量化比选。传统合规审查以条款符合性为主，难以在同一尺度上比较候选方案的"
        "综合竞争力。为此，我们提出 AI Review（人工智能辅助设计评审）框架："
        "将参数化设计候选体映射为五项可观测工程指标，经归一化子得分与加权综合分输出早期方案排序；"
        "并并行构建法规评审通道，在相同五维指标上依据船级社规范与行业阈值独立打分，"
        "通过机队对照检验 AI Review 与法规口径在排序与分维诊断上的一致性。"
        "两通道逻辑解耦、指标同构，使智能体决策既具前瞻性对标能力，又保留法规可追溯性。",
    )

    # ── 2.1 Overview ─────────────────────────────────────────────────────────
    add_heading(doc, "评审框架概览", level=2)
    add_para(
        doc,
        "框架由四个串联模块构成（Fig. 1，示意）："
        "（i）设计候选体表征与物理量提取；"
        "（ii）AI Review 五维打分；"
        "（iii）法规五维打分与定量规则审计；"
        "（iv）机队基准对照与有效性统计。"
        "输入为拓扑优化与几何重建后导出的结构化设计描述；"
        "输出包括综合分、分维得分、规则明细、雷达图、基准位置图及可导出报告。"
        "AI Review 设为主评分通道；法规规则既参与钢耗、疲劳等维度的合成分，"
        "也作为独立 Regulatory Review 供有效性对照，不替代 AI Review 综合分。",
    )

    # ── 2.2 Metric extraction ────────────────────────────────────────────────
    add_heading(doc, "设计候选体表征与指标提取", level=2)
    add_para(
        doc,
        "每个候选体由额定单机容量、设计吃水、壳体代表壁厚、材料密度，"
        "以及支柱几何统计量（平均直径与长度、长细比、变径比、柱间距、"
        "相对等边三角形布局偏差）、轮毂高程、顶盘尺寸与装配体包络尺度等特征描述。"
        "总钢量优先采用外部校核或称重值；否则按几何可用性分层估算："
        "当支柱侧向壳面与顶盘几何完整时，对侧向壳面积积分并乘以壁厚与密度（壳体侧面积法）；"
        "若仅可获得装配体体积，则乘以经标定的壳体系数（体积代理法）。",
    )
    add_equation(doc, "M_steel = ρ · t_w · A_shell + M_plate", "1")
    add_equation(doc, "I_steel = M_steel / P_rated    (t·MW⁻¹)", "2")
    add_para(
        doc,
        "式 (1)–(2) 中，M_steel 为总钢量（t），ρ 为钢材密度（kg·m⁻³），"
        "t_w 为代表壁厚（m），A_shell 为支柱与顶盘侧向壳面积之和（m²），"
        "M_plate 为空心顶盘附加钢量（若适用），P_rated 为额定容量（MW），"
        "I_steel 为钢耗强度。",
    )
    add_para(
        doc,
        "单位造价 C_unit（万元·MW⁻¹）、施工年限 T_const（年）与疲劳设计寿命 T_fat（年）"
        "优先取自项目校核输入；缺失时采用透明工程代理并记入假设清单："
        "C_unit 按 I_steel 相对参考优化方案（I_ref, C_ref）的线性比例估算；"
        "T_const 在 2.0 年基准上叠加总钢量与几何复杂度（变径比）修正；"
        "T_fat 以 25 年设计寿命为基准，并根据直径壁厚比 D/t、壁厚与变径均匀性作疲劳细节代理修正，"
        "修正区间约束在 [18, 30] 年。",
    )
    add_equation(doc, "C_unit = I_steel · (C_ref / I_ref)", "3")
    add_equation(
        doc,
        "T_const = T_0 + min(T_max, M_steel / M_scale) + δ_taper",
        "4",
    )
    add_para(
        doc,
        "式 (3)–(4) 中，I_ref 与 C_ref 为参考方案的钢耗强度与单位造价，"
        "T_0 为基准工期，M_scale 为钢量尺度参数，δ_taper 为变径复杂度修正项。",
    )

    # ── 2.3 AI Review ────────────────────────────────────────────────────────
    add_heading(doc, "AI Review 五维打分模型", level=2)
    add_para(
        doc,
        "AI Review 为每个候选体计算五个子得分 s_i ∈ [0, 100]（i = 1,…,5），"
        "分别对应单机容量、单位兆瓦用钢量、单位造价、施工年限与疲劳寿命。"
        "默认权重 w_i 为 0.15、0.30、0.25、0.15、0.15，满足 Σw_i = 1；"
        "权重可配置但须归一化。综合分 S 定义为加权线性组合：",
    )
    add_equation(doc, "S = Σ_i w_i · s_i", "5")
    add_para(
        doc,
        "综合分映射为等级：A（S ≥ 85）、B（70 ≤ S < 85）、C（60 ≤ S < 70）、D（S < 60）。"
        "各维打分函数均分段线性、单调，并在 [25, 98] 区间内截断，以保证可比性与审计透明度。",
    )

    add_heading(doc, "维度一：单机容量对齐", level=3)
    add_para(
        doc,
        "记候选体额定容量为 P_c，项目目标容量 P*（本研究取 20 MW）。"
        "容量子得分 s_cap 对 |P_c − P*| 分段线性衰减：",
    )
    add_equation(doc, "s_cap = 98，  当 |P_c − P*| ≤ 0.5 MW", "6a")
    add_equation(
        doc,
        "s_cap = 95 − 12(|P_c − P*| − 0.5)，  当 0.5 < |P_c − P*| ≤ 2.0 MW",
        "6b",
    )
    add_equation(
        doc,
        "s_cap = max(45, 72 − 8(|P_c − P*| − 2.0))，  当 |P_c − P*| > 2.0 MW",
        "6c",
    )

    add_heading(doc, "维度二：单位兆瓦用钢量", level=3)
    add_para(
        doc,
        "钢耗强度子得分 s_steel 优先由基准对标规则集给出："
        "对 20 MW 同容量已建成项目样本、300 t·MW⁻¹ 国家示范上限"
        "及相对行业领先方案的裕度等规则进行加权平均。规则不可用时，"
        "在同容量机队样本 {I_k} 中对 I_steel 作经验分位排名（越低越优）：",
    )
    add_equation(doc, "s_steel = Σ_j ω_j · r_j / Σ_j ω_j", "7")
    add_equation(
        doc,
        "s_steel = 100 · (1 − rank(I_steel) / n)，  当 n ≥ 3 且无规则合成分",
        "8",
    )
    add_para(
        doc,
        "式 (7) 中 r_j 为各钢耗相关规则的得分，ω_j 为规则权重；"
        "式 (8) 中 rank 为升序分位排名，n 为同容量样本数。",
    )

    add_heading(doc, "维度三至五：越小越优与越大越优指标", level=3)
    add_para(
        doc,
        "对单位造价与施工年限（越小越优），记实测或代理值为 x，"
        "AI 参考方案值为 x_ref，定义比值 ρ = x / x_ref。"
        "子得分 s(ρ) 采用三段式单调映射：",
    )
    add_equation(doc, "s = 98，  当 ρ ≤ 1.0", "9a")
    add_equation(
        doc,
        "s = 90 − 30(ρ − 1) / (ρ_pass − 1)，  当 1 < ρ ≤ ρ_pass",
        "9b",
    )
    add_equation(
        doc,
        "s = max(25, 60 − 80(ρ − ρ_pass))，  当 ρ > ρ_pass",
        "9c",
    )
    add_para(
        doc,
        "其中 ρ_pass = 1.18。指标缺失时赋予保守基础分 55，并在报告中标注。",
    )
    add_para(
        doc,
        "对疲劳寿命（越大越优），记 L 为候选寿命、L_ref = 25 年为参考，ρ_L = L / L_ref：",
    )
    add_equation(doc, "s = 98，  当 ρ_L ≥ 1", "10a")
    add_equation(
        doc,
        "s = 60 + 38(ρ_L − 0.88) / 0.12，  当 0.88 ≤ ρ_L < 1",
        "10b",
    )
    add_equation(
        doc,
        "s = max(25, 40 + 20 ρ_L / 0.88)，  当 ρ_L < 0.88",
        "10c",
    )

    # ── 2.4 Regulatory Review ────────────────────────────────────────────────
    add_heading(doc, "法规五维评审与定量规则审计", level=2)
    add_para(
        doc,
        "Regulatory Review 在 AI Review 相同的五项指标上独立打分，"
        "但参考基准改为船级社规范与行业阈值，而非 AI 优化参考线。"
        "容量维以 DNV-ST-0437 20 MW 半潜式设计基点及容差带校准；"
        "钢耗维由基准对标规则合成分（含 300 t·MW⁻¹ 国策目标与机队分位规则）；"
        "造价维以国内漂浮式示范 EPC 投资强度控制线（约 3600–4800 万元·MW⁻¹）"
        "映射为式 (9) 型分段函数，其中 x_ref 取合规上限而非 AI 参考价；"
        "工期维以示范工程批复工期惯例（约 2.5–3.5 年）同理映射；"
        "寿命维以 DNV 25 年设计寿命为下限，并与 RP-0286 疲劳细节规则合成分按 0.55:0.45 混合。",
    )
    add_equation(doc, "s_fat^reg = 0.55 · s_life + 0.45 · s_detail", "11")
    add_para(
        doc,
        "并行运行定量规则引擎，对约 26 项可观测指标进行检验，"
        "分属机队基准对标、稳性/水密代理（DNVGL-OS-C301）、"
        "结构布局（DNV-ST-0437）及疲劳/细节代理（DNVGL-RP-0286）四类。"
        "每条规则将指标与阈值比较（分位排名、距上限裕度、区间约束等），"
        "产出规则级得分、通过/警告/失败状态及法规条款溯源。"
        "法规综合分 S_reg 与式 (5) 同构，权重与 AI Review 一致，便于逐维对照。",
    )

    # ── 2.5 Validity ─────────────────────────────────────────────────────────
    add_heading(doc, "机队对照与 AI Review 有效性评估", level=2)
    add_para(
        doc,
        "为检验 AI Review 是否能在同一物理量上复现法规评审的排序逻辑，"
        "我们构建国际与国内已建成及规划漂浮式风电项目机队（n ≥ 11），"
        "对每条记录同时计算 AI 五维得分与法规五维得分。"
        "原始指标缺失的单元格留空（报告中以“—”表示），"
        "不对缺失维度强行插值，以避免虚假相似性。",
    )
    add_para(
        doc,
        "有效性由三类统计量刻画：",
    )
    add_bullet(
        doc,
        "综合分 Spearman 秩相关 ρ_S：衡量 AI 与法规在机队整体排序上的一致性；",
    )
    add_bullet(
        doc,
        "各维及综合分的平均绝对差 MAE = (1/n) Σ |s_i^AI − s_i^reg|；",
    )
    add_bullet(
        doc,
        "高一致占比：双通道综合分均 ≥ 60 且 |S − S_reg| ≤ 15 的项目比例。",
    )
    add_equation(
        doc,
        "ρ_S = 1 − 6 Σ_k d_k² / [n(n² − 1)]",
        "12",
    )
    add_para(
        doc,
        "式 (12) 中 d_k 为第 k 个项目 AI 与法规综合分的秩差。"
        "此外，对每个维度绘制原始性能指数（0–100，由机队内分位归一化）"
        "对 AI/法规得分的散点图，并拟合非线性趋势曲线（二次多项式），"
        "目视检验“物理量改善 → 得分升高”的单调关系是否在双通道间同向。",
    )

    # ── 2.6 Calibration ──────────────────────────────────────────────────────
    add_heading(doc, "校准与不确定性披露", level=2)
    add_para(
        doc,
        "当钢量来自估算而非称重时，对 AI Review 综合分设上限以反映认知不确定性："
        "壳体侧面积法估算时 S ≤ 92，体积代理法时 S ≤ 96；"
        "提供外部校核钢量后解除上限。"
        "验证报告同步输出假设清单、代理来源、机队分位、"
        "相对参考方案偏差、五维雷达图、基准位置回归图及 AI–法规对照表。",
    )

    # ── 2.7 Outputs ──────────────────────────────────────────────────────────
    add_heading(doc, "报告输出与使用边界", level=2)
    add_bullet(doc, "AI Review 综合分 S 与字母等级。")
    add_bullet(doc, "五维子得分、原始物理量及单位。")
    add_bullet(doc, "Regulatory Review 五维得分与规则明细表、条款映射。")
    add_bullet(doc, "机队对照表（原始指标 / AI 得分 / 法规得分）及有效性统计。")
    add_bullet(doc, "基准位置图、五维雷达图、有效性散点图及规则热力图。")
    add_bullet(doc, "可导出的 Markdown、JSON、PNG/PDF 图件及含嵌入图表的 Word 完整报告。")
    add_para(
        doc,
        "说明：AI Review 与 Regulatory Review 得分仅供早期方案比选与智能体决策支持；"
        "正式入级认证与结构实测仍须由持证验船师与审图机构完成。",
        italic=True,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
